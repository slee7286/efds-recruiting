"""Deterministic application artifacts, archives, packets, and verification."""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, cast
from uuid import UUID

from docx import Document
from pypdf import PdfReader
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from quant_recruiting.config import Settings, get_settings
from quant_recruiting.db.models import (
    AITaskRun,
    Application,
    ApplicationAnswer,
    ApplicationArtifact,
    ApplicationQuestion,
    ArtifactProvenance,
    CandidateEvidence,
    CandidateExperience,
    CandidateProfile,
    CoverLetterBlock,
    CoverLetterBlockEvidence,
    CoverLetterBlockSource,
    CVBullet,
    CVBulletEvidence,
    CVVersion,
    ResearchSource,
    ReviewEvent,
    SensitiveField,
)

UTC = getattr(timezone, "UTC", timezone.utc)  # noqa: UP017 - Python 3.10 compatibility

SENSITIVE_FIELDS = {
    "work_authorization",
    "visa_sponsorship",
    "citizenship",
    "nationality",
    "race_ethnicity",
    "gender",
    "disability",
    "veteran_status",
    "criminal_history",
    "salary_expectations",
    "conflicts_of_interest",
    "legal_attestations",
    "relocation_commitment",
}


def sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def sha256_text(value: str) -> str:
    return sha256_bytes(value.encode("utf-8"))


def safe_filename(value: str) -> str:
    clean = re.sub(r"[^A-Za-z0-9._-]+", "_", value).strip("._")
    return clean or "artifact"


def application_root(application_id: UUID, settings: Settings | None = None) -> Path:
    """Return the private local application directory."""
    config = settings or get_settings()
    return config.local_data_dir / "applications" / str(application_id)


def _required_path(value: str | None, label: str) -> Path:
    if not value:
        raise ValueError(f"artifact path missing: {label}")
    return Path(value)


def _next_version(session: Session, application_id: UUID | None, artifact_type: str) -> int:
    value = session.scalar(
        select(func.max(ApplicationArtifact.version)).where(
            ApplicationArtifact.application_id == application_id,
            ApplicationArtifact.artifact_type == artifact_type,
        )
    )
    return int(value or 0) + 1


def _task_provenance(session: Session, task_id: UUID | None) -> tuple[UUID | None, UUID | None]:
    if task_id is None:
        return None, None
    run = session.scalar(
        select(AITaskRun).where(AITaskRun.task_id == task_id).order_by(AITaskRun.run_number.desc())
    )
    return task_id, run.id if run else None


def create_artifact(
    session: Session,
    *,
    application_id: UUID | None,
    artifact_type: str,
    content: bytes,
    source_path: str | None = None,
    rendered_path: str | None = None,
    mime_type: str | None = None,
    source_task_id: UUID | None = None,
    status: str = "draft",
    metadata: dict[str, Any] | None = None,
    provenance: list[tuple[str, UUID, str]] | None = None,
) -> ApplicationArtifact:
    task_id, run_id = _task_provenance(session, source_task_id)
    artifact = ApplicationArtifact(
        application_id=application_id,
        artifact_type=artifact_type,
        version=_next_version(session, application_id, artifact_type),
        status=status,
        source_task_id=task_id,
        source_task_run_id=run_id,
        content_hash=sha256_bytes(content),
        source_path=source_path,
        rendered_path=rendered_path,
        mime_type=mime_type,
        metadata_=metadata or {},
    )
    session.add(artifact)
    session.flush()
    seen_provenance: set[tuple[str, UUID]] = set()
    for entity_type, entity_id, relation in provenance or []:
        if (entity_type, entity_id) in seen_provenance:
            continue
        seen_provenance.add((entity_type, entity_id))
        session.add(
            ArtifactProvenance(
                artifact_id=artifact.id,
                entity_type=entity_type,
                entity_id=entity_id,
                relation=relation,
            )
        )
    session.flush()
    return artifact


def _approved_answers(
    session: Session, application: Application
) -> list[tuple[ApplicationQuestion, ApplicationAnswer]]:
    result: list[tuple[ApplicationQuestion, ApplicationAnswer]] = []
    for question in sorted(application.questions, key=lambda item: item.created_at):
        answer = session.scalar(
            select(ApplicationAnswer)
            .where(
                ApplicationAnswer.application_question_id == question.id,
                ApplicationAnswer.approved.is_(True),
            )
            .order_by(ApplicationAnswer.version.desc())
        )
        if answer:
            result.append((question, answer))
    return result


def _answer_ids(metadata: dict[str, Any], key: str) -> list[UUID]:
    values = metadata.get(key, [])
    return [UUID(str(value)) for value in values]


def _answer_payload(
    application: Application, pairs: list[tuple[ApplicationQuestion, ApplicationAnswer]]
) -> dict[str, Any]:
    return {
        "application_id": str(application.id),
        "company": application.job.company.name,
        "role": application.job.title,
        "location": application.job.location_text,
        "application_url": getattr(application, "application_url", None)
        or getattr(application.job, "job_url", None),
        "questions": [
            {
                "application_question_id": str(question.id),
                "original_question": question.question_text,
                "normalized_category": question.category,
                "required": question.required,
                "max_words": question.max_words,
                "max_characters": question.max_characters,
                "answer": answer.answer_text,
                "word_count": len(answer.answer_text.split()),
                "character_count": len(answer.answer_text),
                "specificity_score": answer.specificity_score,
                "approved": answer.approved,
            }
            for question, answer in pairs
        ],
    }


def _answer_provenance(
    answer: ApplicationAnswer,
) -> list[tuple[str, UUID, str]]:
    metadata = answer.metadata_ or {}
    result: list[tuple[str, UUID, str]] = [("application_answer", answer.id, "renders")]
    result.extend(
        ("candidate_evidence", value, "supports")
        for value in _answer_ids(metadata, "candidate_evidence_ids")
    )
    result.extend(
        ("research_source", value, "supports")
        for value in _answer_ids(metadata, "company_source_ids")
    )
    result.extend(
        ("application_argument", value, "supports")
        for value in _answer_ids(metadata, "argument_ids")
    )
    return result


def render_answers(
    session: Session,
    application: Application,
    settings: Settings,
    *,
    include_provenance_appendix: bool = False,
) -> dict[str, ApplicationArtifact]:
    pairs = _approved_answers(session, application)
    missing = [
        question.id
        for question in application.questions
        if question.required and not any(q.id == question.id for q, _ in pairs)
    ]
    if missing:
        raise ValueError(f"required application questions are not approved: {missing}")
    for question, answer in pairs:
        if question.max_words and len(answer.answer_text.split()) > question.max_words:
            raise ValueError(f"answer exceeds word limit: {question.id}")
        if question.max_characters and len(answer.answer_text) > question.max_characters:
            raise ValueError(f"answer exceeds character limit: {question.id}")
    root = application_root(application.id, settings) / "artifacts" / "answers"
    root.mkdir(parents=True, exist_ok=True)
    payload = _answer_payload(application, pairs)
    json_bytes = (json.dumps(payload, indent=2, ensure_ascii=False) + "\n").encode("utf-8")
    version = _next_version(session, application.id, "written_answers_json")
    json_path = root / f"answers_v{version}.json"
    json_path.write_bytes(json_bytes)
    provenance = [item for _, answer in pairs for item in _answer_provenance(answer)]
    json_artifact = create_artifact(
        session,
        application_id=application.id,
        artifact_type="written_answers_json",
        content=json_bytes,
        rendered_path=str(json_path),
        mime_type="application/json",
        status="validated",
        metadata={"answer_ids": [str(answer.id) for _, answer in pairs]},
        provenance=provenance,
    )
    markdown = _answers_markdown(application, pairs, include_provenance_appendix)
    md_bytes = markdown.encode("utf-8")
    md_path = (
        root / f"answers_v{_next_version(session, application.id, 'written_answers_markdown')}.md"
    )
    md_path.write_bytes(md_bytes)
    md_artifact = create_artifact(
        session,
        application_id=application.id,
        artifact_type="written_answers_markdown",
        content=md_bytes,
        rendered_path=str(md_path),
        mime_type="text/markdown",
        status="validated",
        metadata={"answer_ids": [str(answer.id) for _, answer in pairs]},
        provenance=provenance,
    )
    docx_path = (
        root / f"answers_v{_next_version(session, application.id, 'written_answers_docx')}.docx"
    )
    _write_answers_docx(docx_path, application, pairs, include_provenance_appendix)
    docx_bytes = docx_path.read_bytes()
    docx_artifact = create_artifact(
        session,
        application_id=application.id,
        artifact_type="written_answers_docx",
        content=docx_bytes,
        rendered_path=str(docx_path),
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        status="validated",
        metadata={"answer_ids": [str(answer.id) for _, answer in pairs]},
        provenance=provenance,
    )
    return {"json": json_artifact, "markdown": md_artifact, "docx": docx_artifact}


def _answers_markdown(
    application: Application,
    pairs: list[tuple[ApplicationQuestion, ApplicationAnswer]],
    appendix: bool,
) -> str:
    lines = [
        f"# {application.job.company.name} — {application.job.title}",
        "",
        f"Application ID: `{application.id}`",
        "",
    ]
    for index, (question, answer) in enumerate(pairs, 1):
        lines.extend(
            [
                f"## Question {index}",
                "",
                question.question_text,
                "",
                answer.answer_text,
                "",
                f"Word count: {len(answer.answer_text.split())}",
                f"Character count: {len(answer.answer_text)}",
                "",
            ]
        )
    if appendix:
        lines.extend(["## Internal provenance appendix", ""])
        for question, answer in pairs:
            evidence = answer.metadata_.get("candidate_evidence_ids", [])
            sources = answer.metadata_.get("company_source_ids", [])
            lines.append(
                f"- `{question.id}` -> answer `{answer.id}`; evidence `{evidence}`; "
                f"sources `{sources}`"
            )
    return "\n".join(lines) + "\n"


def _write_answers_docx(
    path: Path,
    application: Application,
    pairs: list[tuple[ApplicationQuestion, ApplicationAnswer]],
    appendix: bool,
) -> None:
    document = Document()
    document.add_heading(f"{application.job.company.name} — {application.job.title}", 0)
    document.add_paragraph(f"Application ID: {application.id}")
    document.add_paragraph(f"Location: {application.job.location_text or 'Not specified'}")
    document.add_paragraph(f"Exported: {datetime.now(UTC).isoformat()}")  # noqa: UP017 - Python 3.10 compatibility
    for index, (question, answer) in enumerate(pairs, 1):
        document.add_heading(f"Question {index}", level=1)
        document.add_paragraph(question.question_text)
        document.add_paragraph(f"Category: {question.category}")
        if question.max_words:
            document.add_paragraph(f"Maximum words: {question.max_words}")
        if question.max_characters:
            document.add_paragraph(f"Maximum characters: {question.max_characters}")
        document.add_heading("Approved answer", level=2)
        document.add_paragraph(answer.answer_text)
        document.add_paragraph(f"Word count: {len(answer.answer_text.split())}")
        document.add_paragraph(f"Character count: {len(answer.answer_text)}")
    if appendix:
        document.add_heading("Internal provenance appendix", level=1)
        for question, answer in pairs:
            evidence = answer.metadata_.get("candidate_evidence_ids", [])
            sources = answer.metadata_.get("company_source_ids", [])
            document.add_paragraph(
                f"Question {question.id}; Answer {answer.id}; Evidence {evidence}; "
                f"Sources {sources}"
            )
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))


def _latex_escape(value: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in value)


def _compile_tex(tex_path: Path, engine: str) -> tuple[Path, str]:
    executable = shutil.which(engine)
    if executable is None:
        raise RuntimeError(f"TeX engine unavailable: {engine}; install MiKTeX or TeX Live")
    result = subprocess.run(
        [
            executable,
            "-interaction=nonstopmode",
            "-halt-on-error",
            "-output-directory",
            str(tex_path.parent),
            str(tex_path),
        ],
        capture_output=True,
        text=True,
        check=False,
    )
    log = (result.stdout or "") + "\n" + (result.stderr or "")
    log_path = tex_path.with_suffix(".log.txt")
    log_path.write_text(log, encoding="utf-8")
    if result.returncode != 0:
        raise RuntimeError(f"TeX compilation failed; see {log_path}")
    pdf_path = tex_path.with_suffix(".pdf")
    if not pdf_path.exists():
        raise RuntimeError("TeX compiler returned success but produced no PDF")
    return pdf_path, log


def _profile(session: Session) -> CandidateProfile:
    profile = session.scalar(select(CandidateProfile).order_by(CandidateProfile.created_at))
    if profile is None:
        raise ValueError("candidate profile is required for deterministic application artifacts")
    return profile


def _approved_cv(session: Session, application: Application) -> tuple[CVVersion, list[CVBullet]]:
    cv = session.scalar(
        select(CVVersion)
        .where(CVVersion.application_id == application.id, CVVersion.approval_status == "approved")
        .order_by(CVVersion.created_at.desc())
    )
    if cv is None:
        raise ValueError("an approved CV version is required")
    bullets = list(
        session.scalars(
            select(CVBullet)
            .where(
                CVBullet.cv_version_id == cv.id,
                CVBullet.approved.is_(True),
                CVBullet.status != "superseded",
            )
            .order_by(CVBullet.order_index)
        )
    )
    if not bullets:
        raise ValueError("approved CV contains no approved bullets")
    for bullet in bullets:
        if not session.scalar(
            select(CVBulletEvidence).where(CVBulletEvidence.bullet_id == bullet.id)
        ):
            raise ValueError(f"CV bullet has no evidence mapping: {bullet.id}")
    return cv, bullets


def _role_template(role_family: str, explicit: str | None) -> str:
    if explicit and explicit not in {"", "default"}:
        return explicit
    lower = role_family.lower()
    if any(term in lower for term in ("investment", "asset", "banking", "equity", "finance")):
        return "finance"
    if any(term in lower for term in ("quant", "trading")):
        return "quant"
    if any(term in lower for term in ("software", "engineering", "machine", "data")):
        return "software"
    return "general"


def render_cv(
    session: Session, application: Application, settings: Settings
) -> dict[str, ApplicationArtifact]:
    profile = _profile(session)
    cv, bullets = _approved_cv(session, application)
    name = profile.preferred_name or profile.legal_name
    if not name:
        raise ValueError("candidate profile name is unresolved")
    template = _role_template(application.job.role_family, cv.template_type)
    template_path = Path("templates") / "cv" / f"{template}.tex"
    if not template_path.exists():
        raise FileNotFoundError(template_path)
    template_text = template_path.read_text(encoding="utf-8")
    if "{{CONTENT}}" not in template_text:
        template_text = (Path("templates") / "cv" / "general.tex").read_text(encoding="utf-8")
    content_lines = [
        f"\\name{{{_latex_escape(name)}}}",
        f"\\role{{{_latex_escape(application.job.title)}}}",
        "\\maketitle",
    ]
    current_section = None
    section_open = False
    for bullet in bullets:
        section = str((bullet.metadata_ or {}).get("section", "Experience"))
        if section != current_section:
            if section_open:
                content_lines.append("\\end{itemize}")
            content_lines.append(f"\\section*{{{_latex_escape(section)}}}")
            content_lines.append("\\begin{itemize}")
            section_open = True
            current_section = section
        evidence_ids = [
            str(row.evidence_id)
            for row in session.scalars(
                select(CVBulletEvidence).where(CVBulletEvidence.bullet_id == bullet.id)
            )
        ]
        content_lines.append(f"% cv_bullet_id: {bullet.id}")
        content_lines.append(f"% evidence_ids: {','.join(evidence_ids)}")
        content_lines.append(f"\\item {_latex_escape(bullet.text)}")
    if section_open:
        content_lines.append("\\end{itemize}")
    body = template_text.replace("{{CONTENT}}", "\n".join(content_lines))
    root = application_root(application.id, settings) / "artifacts" / "cv"
    root.mkdir(parents=True, exist_ok=True)
    source_version = _next_version(session, application.id, "cv_source")
    tex_path = root / f"cv_v{source_version}.tex"
    tex_path.write_bytes(body.encode("utf-8"))
    source_artifact = create_artifact(
        session,
        application_id=application.id,
        artifact_type="cv_source",
        content=body.encode("utf-8"),
        source_path=str(tex_path),
        mime_type="application/x-tex",
        status="validated",
        metadata={"template": template, "cv_version_id": str(cv.id)},
        provenance=[("cv_version", cv.id, "renders")]
        + [("cv_bullet", bullet.id, "renders") for bullet in bullets],
    )
    pdf_path, log = _compile_tex(tex_path, settings.tex_engine)
    reader = PdfReader(str(pdf_path))
    extracted = "\n".join(page.extract_text() or "" for page in reader.pages)
    warnings = []
    if len(reader.pages) > 1:
        warnings.append("CV exceeds the default one-page target")
    if "TODO" in extracted or "{{" in extracted:
        warnings.append("placeholder text detected")
    if name not in extracted:
        warnings.append("candidate name missing from extracted PDF text")
    if warnings:
        raise ValueError("CV validation failed: " + "; ".join(warnings))
    pdf_artifact = create_artifact(
        session,
        application_id=application.id,
        artifact_type="cv_pdf",
        content=pdf_path.read_bytes(),
        rendered_path=str(pdf_path),
        mime_type="application/pdf",
        status="validated",
        metadata={
            "template": template,
            "page_count": len(reader.pages),
            "extracted_text": extracted,
            "compiler_log": log[-4000:],
        },
        provenance=[("cv_version", cv.id, "renders")]
        + [("cv_bullet", bullet.id, "renders") for bullet in bullets],
    )
    return {"source": source_artifact, "pdf": pdf_artifact}


def render_cover_letter(
    session: Session, application: Application, settings: Settings
) -> dict[str, ApplicationArtifact]:
    blocks = list(
        session.scalars(
            select(CoverLetterBlock)
            .where(
                CoverLetterBlock.application_id == application.id,
                CoverLetterBlock.approved.is_(True),
            )
            .order_by(CoverLetterBlock.order_index)
        )
    )
    if not blocks:
        raise ValueError("no approved cover-letter blocks")
    root = application_root(application.id, settings) / "artifacts" / "cover-letter"
    root.mkdir(parents=True, exist_ok=True)
    lines = [f"# {application.job.company.name}", "", f"Re: {application.job.title}", ""] + [
        block.text for block in blocks
    ]
    markdown = "\n\n".join(lines) + "\n"
    md_path = (
        root / f"cover_letter_v{_next_version(session, application.id, 'cover_letter_source')}.md"
    )
    md_path.write_bytes(markdown.encode("utf-8"))
    provenance = [("cover_letter_block", block.id, "renders") for block in blocks]
    for block in blocks:
        provenance.extend(
            ("candidate_evidence", row.evidence_id, "supports")
            for row in session.scalars(
                select(CoverLetterBlockEvidence).where(
                    CoverLetterBlockEvidence.block_id == block.id
                )
            )
        )
        provenance.extend(
            ("research_source", row.source_id, "supports")
            for row in session.scalars(
                select(CoverLetterBlockSource).where(CoverLetterBlockSource.block_id == block.id)
            )
        )
    md_artifact = create_artifact(
        session,
        application_id=application.id,
        artifact_type="cover_letter_source",
        content=markdown.encode("utf-8"),
        source_path=str(md_path),
        rendered_path=str(md_path),
        mime_type="text/markdown",
        status="validated",
        provenance=provenance,
    )
    template_path = Path("templates") / "cover_letter" / "general.tex"
    body = (
        template_path.read_text(encoding="utf-8")
        .replace("{{COMPANY}}", _latex_escape(application.job.company.name))
        .replace("{{ROLE}}", _latex_escape(application.job.title))
        .replace("{{CONTENT}}", "\n\n".join(_latex_escape(block.text) for block in blocks))
    )
    tex_path = root / f"cover_letter_v{md_artifact.version}.tex"
    tex_path.write_bytes(body.encode("utf-8"))
    tex_artifact = create_artifact(
        session,
        application_id=application.id,
        artifact_type="cover_letter_tex",
        content=body.encode("utf-8"),
        source_path=str(tex_path),
        mime_type="application/x-tex",
        status="validated",
        provenance=provenance,
    )
    pdf_path, log = _compile_tex(tex_path, settings.tex_engine)
    reader = PdfReader(str(pdf_path))
    extracted = "\n".join(page.extract_text() or "" for page in reader.pages)
    if application.job.company.name not in extracted:
        raise ValueError("cover letter PDF does not contain company name")
    pdf_artifact = create_artifact(
        session,
        application_id=application.id,
        artifact_type="cover_letter_pdf",
        content=pdf_path.read_bytes(),
        rendered_path=str(pdf_path),
        mime_type="application/pdf",
        status="validated",
        metadata={"page_count": len(reader.pages), "compiler_log": log[-4000:]},
        provenance=provenance,
    )
    return {"markdown": md_artifact, "tex": tex_artifact, "pdf": pdf_artifact}


def _latest_approved_artifact(
    session: Session, application_id: UUID, artifact_type: str
) -> ApplicationArtifact | None:
    return session.scalar(
        select(ApplicationArtifact)
        .where(
            ApplicationArtifact.application_id == application_id,
            ApplicationArtifact.artifact_type == artifact_type,
            ApplicationArtifact.status == "approved",
            ApplicationArtifact.superseded_at.is_(None),
        )
        .order_by(ApplicationArtifact.version.desc())
    )


def _profile_payload(
    session: Session, application: Application
) -> tuple[dict[str, Any], list[str]]:
    profile = session.scalar(select(CandidateProfile).order_by(CandidateProfile.created_at))
    if profile is None:
        return {
            "identity": {},
            "education": [],
            "experience": [],
            "documents": {},
            "questions": [],
            "sensitive_fields": {},
        }, ["candidate_profile"]
    identity = {
        "name": profile.preferred_name or profile.legal_name,
        "email": profile.email,
        "phone": profile.phone,
        "linkedin": profile.linkedin_url,
        "github": profile.github_url,
        "website": profile.website_url,
    }
    unresolved: list[str] = []
    required_sensitive = set(application.job.metadata_.get("required_sensitive_fields", []))
    sensitive: dict[str, str] = {}
    for field in session.scalars(
        select(SensitiveField).where(SensitiveField.profile_id == profile.id)
    ):
        if (
            field.approved
            and field.explicitly_entered
            and field.status == "resolved"
            and field.value is not None
        ):
            sensitive[field.field_name] = field.value
        elif field.field_name in required_sensitive:
            unresolved.append(field.field_name)
    for field in required_sensitive - set(sensitive) - set(unresolved):
        unresolved.append(field)
    experiences = [
        {
            "experience_id": str(experience.id),
            "organization": experience.organization,
            "title": experience.title,
            "description": experience.description,
        }
        for experience in session.scalars(
            select(CandidateExperience)
            .where(CandidateExperience.active.is_(True))
            .order_by(CandidateExperience.start_date.desc().nullslast())
        )
    ]
    return {
        "identity": identity,
        "education": [
            {
                "university": profile.university,
                "degree": profile.degree,
                "subject": profile.degree_subject,
                "graduation_date": profile.graduation_date.isoformat()
                if profile.graduation_date
                else None,
            }
        ],
        "experience": experiences,
        "documents": {},
        "questions": [],
        "sensitive_fields": sensitive,
    }, unresolved


def build_packet(session: Session, application: Application, settings: Settings) -> Path:
    root = application_root(application.id, settings)
    packet = root / "packet"
    packet.mkdir(parents=True, exist_ok=True)
    answer_artifacts_raw = {
        key: _latest_approved_artifact(session, application.id, artifact_type)
        for key, artifact_type in (
            ("json", "written_answers_json"),
            ("markdown", "written_answers_markdown"),
            ("docx", "written_answers_docx"),
        )
    }
    cv_artifacts_raw = {
        key: _latest_approved_artifact(session, application.id, artifact_type)
        for key, artifact_type in (("source", "cv_source"), ("pdf", "cv_pdf"))
    }
    if any(value is None for value in answer_artifacts_raw.values()):
        raise ValueError("approved JSON, Markdown, and DOCX answer artifacts are required")
    if any(value is None for value in cv_artifacts_raw.values()):
        raise ValueError("approved CV source and PDF artifacts are required")
    answer_artifacts: dict[str, ApplicationArtifact] = {
        key: cast(ApplicationArtifact, value) for key, value in answer_artifacts_raw.items()
    }
    cv_artifacts: dict[str, ApplicationArtifact] = {
        key: cast(ApplicationArtifact, value) for key, value in cv_artifacts_raw.items()
    }
    optional_cover: dict[str, ApplicationArtifact] | None = None
    if application.cover_letter_requirement != "not_required":
        optional_cover_raw = {
            key: _latest_approved_artifact(session, application.id, artifact_type)
            for key, artifact_type in (
                ("markdown", "cover_letter_source"),
                ("tex", "cover_letter_tex"),
                ("pdf", "cover_letter_pdf"),
            )
        }
        if any(value is None for value in optional_cover_raw.values()):
            if application.cover_letter_requirement == "required":
                raise ValueError("approved cover-letter artifacts are required")
        else:
            optional_cover = {
                key: cast(ApplicationArtifact, value) for key, value in optional_cover_raw.items()
            }
    files: dict[str, str] = {}
    artifact_files: dict[str, tuple[ApplicationArtifact, str]] = {
        "answers_json": (answer_artifacts["json"], "answers.json"),
        "answers_markdown": (answer_artifacts["markdown"], "answers.md"),
        "answers_docx": (answer_artifacts["docx"], "answers.docx"),
        "cv_tex": (cv_artifacts["source"], "cv.tex"),
        "cv_pdf": (cv_artifacts["pdf"], "cv.pdf"),
    }
    if optional_cover:
        artifact_files.update(
            {
                "cover_letter_md": (optional_cover["markdown"], "cover-letter.md"),
                "cover_letter_tex": (optional_cover["tex"], "cover-letter.tex"),
                "cover_letter_pdf": (optional_cover["pdf"], "cover-letter.pdf"),
            }
        )
    for key, (artifact, filename) in artifact_files.items():
        source = artifact.rendered_path or artifact.source_path
        if source is None:
            raise ValueError(f"artifact has no file path: {artifact.id}")
        destination = packet / filename
        shutil.copyfile(source, destination)
        files[key] = str(destination)
    payload, unresolved = _profile_payload(session, application)
    payload["application_id"] = str(application.id)
    payload["job"] = {
        "company": application.job.company.name,
        "role": application.job.title,
        "application_url": application.application_url or application.job.job_url,
    }
    answer_json_path = _required_path(answer_artifacts["json"].rendered_path, "answers JSON")
    answer_json = json.loads(answer_json_path.read_text(encoding="utf-8"))
    payload["questions"] = answer_json["questions"]
    documents = payload["documents"]
    documents["cv"] = files.get("cv_pdf")
    documents["cover_letter"] = files.get("cover_letter_pdf")
    payload["documents"] = documents
    payload["unresolved_fields"] = unresolved
    payload_path = packet / "application_form_payload.json"
    payload_path.write_text(
        json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
    )
    summary = application_summary(application, unresolved, session)
    (packet / "application-summary.md").write_text(summary, encoding="utf-8")
    job_markdown = "\n".join(
        [
            "# Job",
            "",
            f"- Company: {application.job.company.name}",
            f"- Role: {application.job.title}",
            f"- Role family: {application.job.role_family}",
            f"- Location: {application.job.location_text or 'Not specified'}",
            f"- Job URL: {application.job.job_url}",
            "",
            "## Description",
            "",
            application.job.description_text or "No description captured.",
            "",
            "## Requirements",
            "",
            application.job.requirements_text or "No requirements captured.",
            "",
        ]
    )
    (packet / "job.md").write_bytes(job_markdown.encode("utf-8"))
    sources = list(
        session.scalars(
            select(ResearchSource)
            .where(ResearchSource.company_id == application.job.company_id)
            .order_by(ResearchSource.retrieved_at.desc())
        )
    )
    company_markdown = "\n".join(
        [
            "# Company brief",
            "",
            f"- Company: {application.job.company.name}",
            f"- Primary domain: {application.job.company.primary_domain or 'Not configured'}",
            f"- Research sources available: {len(sources)}",
            "",
            "This is a deterministic index. It contains no generated conclusions.",
            "",
        ]
    )
    (packet / "company-brief.md").write_bytes(company_markdown.encode("utf-8"))
    interview_markdown = "\n".join(
        [
            "# Interview preparation",
            "",
            f"- Company: {application.job.company.name}",
            f"- Role: {application.job.title}",
            "",
            "Use the approved interview-preparation records and cited research sources "
            "from the application workspace. No process claims are inferred here.",
            "",
        ]
    )
    (packet / "interview-prep.md").write_bytes(interview_markdown.encode("utf-8"))
    sources_payload = [
        {
            "source_id": str(source.id),
            "title": source.title,
            "url": source.url,
            "source_type": source.source_type,
            "source_quality": source.source_quality,
            "retrieved_at": source.retrieved_at.isoformat(),
        }
        for source in sources
    ]
    sources_path = packet / "sources.json"
    sources_path.write_bytes(
        (json.dumps(sources_payload, indent=2, ensure_ascii=False) + "\n").encode("utf-8")
    )
    manifest = {
        "application_id": str(application.id),
        "job_id": str(application.job.id),
        "company_id": str(application.job.company_id),
        "role": application.job.title,
        "location": application.job.location_text,
        "application_url": application.application_url or application.job.job_url,
        "artifact_ids": [
            str(artifact.id)
            for artifact in [
                *answer_artifacts.values(),
                *cv_artifacts.values(),
                *(list(optional_cover.values()) if optional_cover else []),
            ]
        ],
        "artifact_file_ids": {
            key: str(artifact.id) for key, (artifact, _filename) in artifact_files.items()
        },
        "files": files,
        "supporting_files": {
            "job": str(packet / "job.md"),
            "company_brief": str(packet / "company-brief.md"),
            "interview_prep": str(packet / "interview-prep.md"),
            "sources": str(sources_path),
        },
        "source_ids": [str(source.id) for source in sources],
        "payload": str(payload_path),
        "unresolved_fields": unresolved,
        "generated_at": datetime.now(UTC).isoformat(),
        "packet_version": 1,
    }
    (packet / "manifest.json").write_text(
        json.dumps(manifest, indent=2, sort_keys=True) + "\n", encoding="utf-8"
    )
    return packet


def application_summary(application: Application, unresolved: list[str], session: Session) -> str:
    gaps = list(
        session.scalars(
            select(ReviewEvent).where(
                ReviewEvent.application_id == application.id, ReviewEvent.action == "gap"
            )
        )
    )
    lines = [
        "# Application summary",
        "",
        f"- Company: {application.job.company.name}",
        f"- Role: {application.job.title}",
        f"- Role family: {application.job.role_family}",
        f"- Location: {application.job.location_text or 'Not specified'}",
        f"- Job URL: {application.job.job_url}",
        f"- Application URL: {application.application_url or application.job.job_url}",
        f"- Status: {application.status}",
        f"- Cover letter: {application.cover_letter_requirement}",
        f"- Unresolved sensitive fields: {', '.join(unresolved) or 'none'}",
        f"- Review events: {len(gaps)}",
        "",
        "## Review checklist",
        "",
        "- [ ] Review CV",
        "- [ ] Review written answers",
        "- [ ] Review cover letter if applicable",
        "- [ ] Verify application form payload",
        "- [ ] Confirm no blocking gaps",
    ]
    return "\n".join(lines) + "\n"


def verify_packet(session: Session, application: Application) -> dict[str, Any]:
    packet = application_root(application.id) / "packet"
    errors: list[str] = []
    manifest_path = packet / "manifest.json"
    payload_path = packet / "application_form_payload.json"
    if not manifest_path.exists() or not payload_path.exists():
        return {"status": "invalid", "errors": ["packet manifest or payload missing"]}
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    for key, filename in manifest.get("files", {}).items():
        if not Path(filename).exists():
            errors.append(f"missing packet file: {filename}")
        artifact_id = manifest.get("artifact_file_ids", {}).get(key)
        if artifact_id:
            artifact = session.get(ApplicationArtifact, UUID(artifact_id))
            if artifact is None or artifact.application_id != application.id:
                errors.append(f"artifact does not belong to application: {artifact_id}")
            elif artifact.status != "approved":
                errors.append(f"artifact is not approved: {artifact_id}")
            elif (
                Path(filename).exists()
                and sha256_bytes(Path(filename).read_bytes()) != artifact.content_hash
            ):
                errors.append(f"artifact hash mismatch: {artifact_id}")
    for key, filename in manifest.get("supporting_files", {}).items():
        if not Path(filename).exists():
            errors.append(f"missing supporting packet file ({key}): {filename}")
    answers_path = packet / "answers.json"
    if not answers_path.exists():
        errors.append("answers.json missing")
    else:
        actual = json.loads(answers_path.read_text(encoding="utf-8"))
        expected = _answer_payload(application, _approved_answers(session, application))
        if actual != expected:
            errors.append("answers.json does not equal approved answers")
        markdown_path = packet / "answers.md"
        if not markdown_path.exists():
            errors.append("answers.md missing")
        else:
            markdown_text = markdown_path.read_text(encoding="utf-8")
            for question in expected["questions"]:
                if question["answer"] not in markdown_text:
                    errors.append(
                        "answer missing exactly from Markdown: "
                        f"{question['application_question_id']}"
                    )
        docx_path = packet / "answers.docx"
        if not docx_path.exists():
            errors.append("answers.docx missing")
        else:
            text = "\n".join(paragraph.text for paragraph in Document(str(docx_path)).paragraphs)
            for question in expected["questions"]:
                if question["answer"] not in text:
                    errors.append(
                        f"answer missing exactly from DOCX: {question['application_question_id']}"
                    )
    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    if payload.get("unresolved_fields") is None:
        errors.append("payload unresolved_fields missing")
    return {"status": "valid" if not errors else "invalid", "errors": errors, "packet": str(packet)}


def approve_artifact(session: Session, artifact_id: UUID, reviewer: str) -> ApplicationArtifact:
    artifact = session.get(ApplicationArtifact, artifact_id)
    if artifact is None:
        raise LookupError(f"unknown artifact: {artifact_id}")
    if artifact.status != "validated":
        raise ValueError("only validated artifacts can be approved")
    for previous in session.scalars(
        select(ApplicationArtifact).where(
            ApplicationArtifact.application_id == artifact.application_id,
            ApplicationArtifact.artifact_type == artifact.artifact_type,
            ApplicationArtifact.status == "approved",
            ApplicationArtifact.id != artifact.id,
            ApplicationArtifact.superseded_at.is_(None),
        )
    ):
        previous.status = "superseded"
        previous.superseded_at = datetime.now(UTC)
    artifact.status = "approved"
    artifact.approved_at = datetime.now(UTC)  # noqa: UP017 - Python 3.10 compatibility
    session.add(
        ReviewEvent(
            entity_type="application_artifact",
            entity_id=artifact.id,
            application_id=artifact.application_id,
            action="approval",
            previous_status="validated",
            new_status="approved",
            occurred_at=datetime.now(UTC),
            metadata_={"reviewer": reviewer},
        )
    )
    session.flush()
    return artifact


def approve_answer(session: Session, answer_id: UUID, reviewer: str) -> ApplicationAnswer:
    answer = session.get(ApplicationAnswer, answer_id)
    if answer is None:
        raise LookupError(f"unknown application answer: {answer_id}")
    answer.approved = True
    answer.metadata_["provenance_review_required"] = False
    question = answer.question
    session.add(
        ReviewEvent(
            entity_type="application_answer",
            entity_id=answer.id,
            application_id=question.application_id,
            action="approval",
            previous_status="draft",
            new_status="approved",
            occurred_at=datetime.now(UTC),
            metadata_={"reviewer": reviewer},
        )
    )
    session.flush()
    return answer


def edit_answer(session: Session, answer_id: UUID, text: str, reviewer: str) -> ApplicationAnswer:
    answer = session.get(ApplicationAnswer, answer_id)
    if answer is None:
        raise LookupError(f"unknown application answer: {answer_id}")
    if not text.strip():
        raise ValueError("edited answer cannot be empty")
    latest = session.scalar(
        select(func.max(ApplicationAnswer.version)).where(
            ApplicationAnswer.application_question_id == answer.application_question_id
        )
    )
    metadata = dict(answer.metadata_ or {})
    metadata["provenance_review_required"] = True
    metadata["edited_from_answer_id"] = str(answer.id)
    edited = ApplicationAnswer(
        application_question_id=answer.application_question_id,
        answer_text=text,
        version=int(latest or 0) + 1,
        approved=False,
        specificity_score=answer.specificity_score,
        generated_at=datetime.now(UTC),
        metadata_=metadata,
    )
    session.add(edited)
    session.flush()
    session.add(
        ReviewEvent(
            entity_type="application_answer",
            entity_id=edited.id,
            application_id=answer.question.application_id,
            action="edit",
            previous_status="approved" if answer.approved else "draft",
            new_status="draft",
            occurred_at=datetime.now(UTC),
            metadata_={"reviewer": reviewer, "source_answer_id": str(answer.id)},
        )
    )
    session.flush()
    return edited


def application_readiness(session: Session, application: Application) -> dict[str, Any]:
    cv = _latest_approved_artifact(session, application.id, "cv_pdf")
    answers = _approved_answers(session, application)
    required_answers = [question for question in application.questions if question.required]
    answer_ok = len(answers) == len(required_answers)
    docx = _latest_approved_artifact(session, application.id, "written_answers_docx")
    json_artifact = _latest_approved_artifact(session, application.id, "written_answers_json")
    cover = _latest_approved_artifact(session, application.id, "cover_letter_pdf")
    packet_result = (
        verify_packet(session, application)
        if (application_root(application.id) / "packet").exists()
        else {"status": "not_built", "errors": []}
    )
    payload, unresolved = _profile_payload(session, application)
    blocking = list(unresolved)
    if application.job.status == "closed":
        blocking.append("job is closed")
    if not (application.application_url or application.job.job_url):
        blocking.append("application URL")
    identity = payload.get("identity", {})
    if not identity.get("name"):
        blocking.append("candidate name")
    if not identity.get("email"):
        blocking.append("candidate email")
    if cv is None:
        blocking.append("approved valid CV")
    if not answer_ok:
        blocking.append("all required answers approved")
    if docx is None:
        blocking.append("approved answers DOCX")
    if json_artifact is None:
        blocking.append("approved answers JSON")
    if application.cover_letter_requirement == "required" and cover is None:
        blocking.append("approved cover letter")
    if packet_result["status"] != "valid":
        blocking.append("packet verification")
    return {
        "application_id": str(application.id),
        "cv": bool(cv),
        "answers": f"{len(answers)}/{len(required_answers)}",
        "answers_docx": bool(docx),
        "answers_json": bool(json_artifact),
        "cover_letter": bool(cover)
        if application.cover_letter_requirement != "not_required"
        else "N/A",
        "payload": bool(payload),
        "unresolved_fields": unresolved,
        "blocking_gaps": blocking,
        "packet": packet_result["status"],
        "status": "READY TO APPLY" if not blocking else "NOT READY",
    }


def export_candidate_profile(session: Session) -> Path:
    profile = session.scalar(select(CandidateProfile).order_by(CandidateProfile.created_at))
    if profile is None:
        raise ValueError("candidate profile not found")
    path = get_settings().local_data_dir / "exports" / "candidate-profile.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    lines = [
        "# Candidate profile",
        "",
        f"- Name: {profile.preferred_name or profile.legal_name or 'Unresolved'}",
        f"- Email: {profile.email or 'Unresolved'}",
        f"- University: {profile.university or 'Unresolved'}",
        "",
        "## Experiences",
        "",
    ]
    for experience in session.scalars(
        select(CandidateExperience)
        .where(CandidateExperience.active.is_(True))
        .order_by(CandidateExperience.start_date.desc().nullslast())
    ):
        lines.extend(
            [
                f"### {experience.title} — {experience.organization}",
                "",
                experience.description or "",
                "",
            ]
        )
    lines.extend(["## Approved evidence", ""])
    for evidence in session.scalars(
        select(CandidateEvidence).where(CandidateEvidence.approved_for_application.is_(True))
    ):
        lines.append(f"- {evidence.statement}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def diff_cv(session: Session, first_id: UUID, second_id: UUID) -> dict[str, Any]:
    first = session.get(ApplicationArtifact, first_id)
    second = session.get(ApplicationArtifact, second_id)
    if first is None or second is None:
        raise LookupError("both CV artifact IDs must exist")
    if first.artifact_type != "cv_pdf" or second.artifact_type != "cv_pdf":
        raise ValueError("CV diff requires cv_pdf artifacts")
    first_bullets = {
        row.entity_id
        for row in session.scalars(
            select(ArtifactProvenance).where(
                ArtifactProvenance.artifact_id == first.id,
                ArtifactProvenance.entity_type == "cv_bullet",
            )
        )
    }
    second_bullets = {
        row.entity_id
        for row in session.scalars(
            select(ArtifactProvenance).where(
                ArtifactProvenance.artifact_id == second.id,
                ArtifactProvenance.entity_type == "cv_bullet",
            )
        )
    }
    return {
        "artifact_a": str(first.id),
        "artifact_b": str(second.id),
        "template_a": first.metadata_.get("template"),
        "template_b": second.metadata_.get("template"),
        "bullets_added": [str(value) for value in sorted(second_bullets - first_bullets, key=str)],
        "bullets_removed": [
            str(value) for value in sorted(first_bullets - second_bullets, key=str)
        ],
        "content_changed": first.content_hash != second.content_hash,
        "page_count_a": first.metadata_.get("page_count"),
        "page_count_b": second.metadata_.get("page_count"),
    }
