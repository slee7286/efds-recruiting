from pathlib import Path
from types import SimpleNamespace
from uuid import uuid4

from docx import Document

from quant_recruiting.artifacts import _answer_payload, _latex_escape, _write_answers_docx
from quant_recruiting.db.base import Base


def test_v6_tables_are_present() -> None:
    expected = {
        "candidate_profiles",
        "candidate_sensitive_fields",
        "candidate_cv_sections",
        "candidate_cv_entries",
        "application_artifacts",
        "artifact_provenance",
        "application_answer_evidence",
        "application_answer_sources",
        "cover_letter_blocks",
        "cover_letter_block_evidence",
        "cover_letter_block_sources",
        "review_events",
        "browser_fill_runs",
        "browser_field_mappings",
    }
    assert expected <= set(Base.metadata.tables)


def _context() -> tuple[SimpleNamespace, SimpleNamespace, SimpleNamespace]:
    question = SimpleNamespace(
        id=uuid4(),
        question_text="Why this firm? Café — 東京",
        category="why_company",
        max_words=200,
        max_characters=None,
        required=True,
    )
    answer = SimpleNamespace(
        id=uuid4(),
        answer_text="I value the firm’s research culture — especially its public work.",
        specificity_score=5,
        approved=True,
        metadata_={"candidate_evidence_ids": [], "company_source_ids": []},
    )
    company = SimpleNamespace(name="Example Firm")
    job = SimpleNamespace(company=company, title="Software Engineer Intern", location_text="London")
    application = SimpleNamespace(id=uuid4(), job=job)
    return application, question, answer


def test_answer_payload_preserves_exact_text_and_counts() -> None:
    application, question, answer = _context()
    payload = _answer_payload(application, [(question, answer)])
    item = payload["questions"][0]
    assert item["original_question"] == question.question_text
    assert item["answer"] == answer.answer_text
    assert item["character_count"] == len(answer.answer_text)
    assert item["word_count"] == len(answer.answer_text.split())


def test_docx_archive_preserves_unicode_question_and_answer(tmp_path: Path) -> None:
    application, question, answer = _context()
    path = tmp_path / "answers.docx"
    _write_answers_docx(path, application, [(question, answer)], False)
    text = "\n".join(paragraph.text for paragraph in Document(str(path)).paragraphs)
    assert question.question_text in text
    assert answer.answer_text in text
    assert "Word count:" in text


def test_latex_escape_handles_special_characters() -> None:
    escaped = _latex_escape(r"A&B_50% $x$")
    assert r"A\&B\_50\% \$x\$" == escaped
